"""
File ingestion pipeline for processing documents and creating vector embeddings.
Handles file type detection, text extraction, chunking, and vector store integration.
"""

from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import hashlib
import mimetypes
from datetime import datetime
import re

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from .logging_config import get_logger
from .vector_manager import VectorManager

logger = get_logger(__name__)


class FileIngestor:
    """
    Handles file ingestion, processing, and vector store integration.
    Supports multiple file formats with configurable chunking and embedding strategies.
    """

    # Supported file types and their MIME types
    SUPPORTED_TYPES = {
        'text/plain': ['.txt', '.md', '.py', '.js', '.html', '.css', '.json', '.xml', '.yaml', '.yml'],
        'application/pdf': ['.pdf'],
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': ['.docx'],
        'application/msword': ['.doc'],
        'text/csv': ['.csv'],
        'application/json': ['.json'],
        'text/markdown': ['.md'],
    }

    def __init__(self,
                 vector_store_path: str = "./data/vector_store",
                 embedding_provider: str = "openai",
                 embedding_model: str = "text-embedding-3-small",
                 api_key: Optional[str] = None,
                 vector_provider: str = "faiss"):
        """
        Initialize the file ingestor.

        Args:
            vector_store_path: Path to store vector database
            embedding_provider: Embedding provider ('openai' or 'openrouter')
            embedding_model: Embedding model to use
            api_key: API key for the embedding provider
            vector_provider: Vector store provider ('faiss' or 'chroma')
        """
        # Initialize text splitter with sensible defaults
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", ". ", " ", ""]
        )

        # Initialize vector manager
        self.vector_manager = VectorManager(
            provider=vector_provider,
            store_path=vector_store_path,
            embedding_provider=embedding_provider,
            embedding_model=embedding_model,
            api_key=api_key
        )

        logger.info(f"FileIngestor initialized with {vector_provider} vector store and {embedding_provider} embeddings at {vector_store_path}")

    def detect_file_type(self, file_path: str) -> Tuple[str, bool]:
        """
        Detect file type and check if it's supported.

        Args:
            file_path: Path to the file

        Returns:
            Tuple of (mime_type, is_supported)
        """
        mime_type, _ = mimetypes.guess_type(file_path)

        if mime_type and mime_type in self.SUPPORTED_TYPES:
            return mime_type, True

        # Check file extension as fallback
        file_ext = Path(file_path).suffix.lower()
        for mime, extensions in self.SUPPORTED_TYPES.items():
            if file_ext in extensions:
                return mime, True

        return mime_type or "unknown", False

    def extract_text(self, file_path: str, mime_type: str) -> str:
        """
        Extract text content from file based on its type.

        Args:
            file_path: Path to the file
            mime_type: MIME type of the file

        Returns:
            Extracted text content

        Raises:
            ValueError: If file type is not supported or extraction fails
        """
        file_path_obj = Path(file_path)

        if not file_path_obj.exists():
            raise ValueError(f"File does not exist: {file_path}")

        try:
            if mime_type == 'text/plain' or file_path_obj.suffix.lower() in ['.txt', '.md', '.py', '.js', '.html', '.css', '.json', '.xml', '.yaml', '.yml']:
                return self._extract_text_file(file_path)
            elif mime_type == 'application/pdf':
                return self._extract_pdf_file(file_path)
            elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
                return self._extract_docx_file(file_path)
            else:
                raise ValueError(f"Unsupported file type: {mime_type}")

        except Exception as e:
            logger.error(f"Failed to extract text from {file_path}: {e}")
            raise ValueError(f"Text extraction failed: {str(e)}")

    def _extract_text_file(self, file_path: str) -> str:
        """Extract text from plain text files."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    def _extract_pdf_file(self, file_path: str) -> str:
        """Extract text from PDF files."""
        try:
            import PyPDF2
        except ImportError:
            raise ValueError("PyPDF2 not installed. Install with: pip install PyPDF2")

        text = ""
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"

        return text

    def _extract_docx_file(self, file_path: str) -> str:
        """Extract text from Word documents."""
        try:
            from docx import Document
        except ImportError:
            raise ValueError("python-docx not installed. Install with: pip install python-docx")

        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        return text

    def chunk_text(self, text: str, metadata: Dict[str, Any]) -> List[Document]:
        """
        Split text into chunks with metadata.

        Args:
            text: Text content to chunk
            metadata: Metadata to attach to each chunk

        Returns:
            List of Document objects
        """
        # Create base document
        doc = Document(page_content=text, metadata=metadata)

        # Split into chunks
        chunks = self.text_splitter.split_documents([doc])

        # Add chunk-specific metadata
        for i, chunk in enumerate(chunks):
            chunk.metadata.update({
                'chunk_id': i,
                'total_chunks': len(chunks),
                'chunk_size': len(chunk.page_content)
            })

        logger.info(f"Split document into {len(chunks)} chunks")
        return chunks

    def generate_file_metadata(self, file_path: str, mime_type: str) -> Dict[str, Any]:
        """
        Generate metadata for a file.

        Args:
            file_path: Path to the file
            mime_type: MIME type of the file

        Returns:
            Metadata dictionary
        """
        file_path_obj = Path(file_path)
        stat = file_path_obj.stat()

        # Generate file hash for deduplication
        file_hash = self._calculate_file_hash(file_path)

        metadata = {
            'file_path': str(file_path),
            'file_name': file_path_obj.name,
            'file_size': stat.st_size,
            'file_modified': datetime.fromtimestamp(stat.st_mtime).isoformat(),
            'mime_type': mime_type,
            'file_hash': file_hash,
            'ingestion_time': datetime.now().isoformat(),
        }

        return metadata

    def _calculate_file_hash(self, file_path: str) -> str:
        """Calculate SHA256 hash of file content."""
        hash_sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()

    def ingest_file(self, file_path: str) -> Dict[str, Any]:
        """
        Ingest a single file: extract text, chunk, and add to vector store.

        Args:
            file_path: Path to the file to ingest

        Returns:
            Ingestion result with metadata

        Raises:
            ValueError: If file type is not supported or processing fails
        """
        # Input validation
        if not file_path or not isinstance(file_path, str):
            raise ValueError("File path must be a non-empty string")

        file_path = file_path.strip()
        if not file_path:
            raise ValueError("File path cannot be empty or whitespace")

        # Convert to Path object and resolve
        try:
            file_path_obj = Path(file_path).resolve()
        except Exception as e:
            raise ValueError(f"Invalid file path: {e}")

        # Check if file exists
        if not file_path_obj.exists():
            raise ValueError(f"File does not exist: {file_path}")

        # Check if it's actually a file
        if not file_path_obj.is_file():
            raise ValueError(f"Path is not a file: {file_path}")

        # Check file size (reasonable limit: 50MB)
        max_size = 50 * 1024 * 1024  # 50MB
        file_size = file_path_obj.stat().st_size
        if file_size > max_size:
            raise ValueError(f"File too large: {file_size} bytes (max {max_size} bytes)")

        if file_size == 0:
            raise ValueError(f"File is empty: {file_path}")

        logger.info(f"Starting ingestion of file: {file_path} ({file_size} bytes)")

        # Detect file type
        mime_type, is_supported = self.detect_file_type(str(file_path_obj))
        if not is_supported:
            raise ValueError(f"Unsupported file type: {mime_type} for file {file_path}")

        # Generate metadata
        metadata = self.generate_file_metadata(str(file_path_obj), mime_type)

        try:
            # Check if file already exists (by hash) - this will trigger embeddings initialization
            existing_files = self.vector_manager.get_files_info()
            if any(f.get('file_hash') == metadata['file_hash'] for f in existing_files):
                logger.info(f"File {file_path} already ingested (hash: {metadata['file_hash']})")
                return {
                    'status': 'already_exists',
                    'file_path': file_path,
                    'file_hash': metadata['file_hash'],
                    'chunks_added': 0
                }

            # Extract text
            text = self.extract_text(file_path, mime_type)

            if not text.strip():
                raise ValueError(f"No text content extracted from {file_path}")

            # Chunk text
            chunks = self.chunk_text(text, metadata)

            # Add to vector store
            self.vector_manager.add_documents(chunks)

            result = {
                'status': 'success',
                'file_path': file_path,
                'file_hash': metadata['file_hash'],
                'chunks_added': len(chunks),
                'total_size': len(text),
                'mime_type': mime_type
            }

            logger.info(f"Successfully ingested file: {file_path} ({len(chunks)} chunks)")
            return result

        except ValueError as e:
            if "OpenAI API key" in str(e):
                logger.error(f"Cannot ingest file: {e}")
                return {
                    'status': 'error',
                    'file_path': file_path,
                    'error': 'OpenAI API key not configured. Please configure your API key in settings.',
                    'chunks_added': 0
                }
            else:
                raise

    def search_similar(self, query: str, k: int = 5, score_threshold: float = 0.0) -> List[Dict[str, Any]]:
        """
        Search for similar documents in the vector store.

        Args:
            query: Search query
            k: Number of results to return
            score_threshold: Minimum similarity score (0.0 to 1.0)

        Returns:
            List of search results with metadata and scores
        """
        return self.vector_manager.search_similar(query, k, score_threshold)

    def hybrid_search(self, query: str, k: int = 5, semantic_weight: float = 0.7, keyword_weight: float = 0.3) -> List[Dict[str, Any]]:
        """
        Perform hybrid search combining semantic and keyword-based retrieval.

        Args:
            query: Search query
            k: Number of results to return
            semantic_weight: Weight for semantic similarity (0.0 to 1.0)
            keyword_weight: Weight for keyword matching (0.0 to 1.0)

        Returns:
            List of search results with combined scores
        """
        return self.vector_manager.hybrid_search(query, k, semantic_weight, keyword_weight)

    def get_ingested_files(self) -> List[Dict[str, Any]]:
        """
        Get list of all ingested files with metadata.

        Returns:
            List of file metadata
        """
        return self.vector_manager.get_files_info()

    def delete_file(self, file_path: str) -> bool:
        """
        Delete a file from the vector store by path.

        Args:
            file_path: Path of file to delete

        Returns:
            True if deletion was successful
        """
        try:
            # Delete by file path filter
            deleted_count = self.vector_manager.delete_documents({'file_path': file_path})
            return deleted_count > 0
        except Exception as e:
            logger.error(f"Failed to delete file {file_path}: {e}")
            return False

    def rebuild_index(self):
        """Rebuild the vector store index. Useful after bulk operations."""
        try:
            self.vector_manager.rebuild_index()
        except Exception as e:
            logger.error(f"Failed to rebuild index: {e}")
            raise
